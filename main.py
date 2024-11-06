import os
from docx import Document
from docx.shared import Inches, Pt
from pypdf import PdfMerger, PdfReader
from pyrogram import Client, filters
from pyrogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from pyrogram.errors import PeerFloodError, UserPrivacyRestrictedError
from io import BytesIO
from pypdf import PdfReader, PdfWriter

# -------------------------------------------------------------------------------------------------------------------- #
# Credenciales del Bot (No se recomienda guardarlas en el script principal)
# -------------------------------------------------------------------------------------------------------------------- #
api_id = 24288670 # Reemplaza con tu API ID
api_hash = "81c58005802498656d6b689dae1edacc" # Reemplaza con tu API HASH
bot_token = "7507770865:AAFDQ0Lbuo5Ca-mTnqSa-dK_UJENs5B2v1Q" # Reemplaza con tu Bot Token
# -------------------------------------------------------------------------------------------------------------------- #

# -------------------------------------------------------------------------------------------------------------------- #
# Iniciamos el Bot
# -------------------------------------------------------------------------------------------------------------------- #
app = Client("my_bot", api_id=api_id, api_hash=api_hash, bot_token=bot_token)
# -------------------------------------------------------------------------------------------------------------------- #

# -------------------------------------------------------------------------------------------------------------------- #
# Variables Globales
# -------------------------------------------------------------------------------------------------------------------- #
documents = []
current_file_type = None
customization_options = {
    "nombre_archivo": "archivo_combinado",
    "formato": "pdf",
    "tamaño_fuente": 12,
}
# -------------------------------------------------------------------------------------------------------------------- #

# -------------------------------------------------------------------------------------------------------------------- #
# Funciones para el Manejo de Archivos
# -------------------------------------------------------------------------------------------------------------------- #
def leer_pdf(archivo_pdf):
    """Lee un archivo PDF y devuelve su contenido como texto."""
    with open(archivo_pdf, 'rb') as f:
        parser = PDFParser(f)
        document = PDFDocument(parser)
        content = ''
        for page in document.get_pages():
            content += PDFPageContent.create_content(page).get_text()
    return content

def combinar_archivos(archivos, formato_salida, nombre_salida, opciones):
    """Combina archivos en el formato de salida especificado."""
    if formato_salida == 'pdf':
        merger = PdfMerger()
        for archivo in archivos:
            merger.append(archivo)
        merger.write(f"{nombre_salida}.pdf")
        merger.close()

    elif formato_salida == 'word':
        doc = Document()
        for archivo in archivos:
            if archivo.endswith('.docx') or archivo.endswith('.doc'):
                doc_temp = Document(archivo)
                for parrafo in doc_temp.paragraphs:
                    doc.add_paragraph(parrafo.text)
            elif archivo.endswith('.txt'):
                with open(archivo, 'r') as f:
                    doc.add_paragraph(f.read())

        # Ajusta tamaño de fuente (opcional)
        if 'tamaño_fuente' in opciones:
            for parrafo in doc.paragraphs:
                parrafo.style.font.size = Pt(opciones['tamaño_fuente'])

        doc.save(f"{nombre_salida}.docx")

def convertir_pdf_a_word(archivo_pdf, nombre_salida, tamaño_fuente):
    """Convierte un archivo PDF a Word."""
    reader = PdfReader(archivo_pdf)
    num_pages = len(reader.pages)
    doc = Document()
    for page_num in range(num_pages):
        page = reader.pages[page_num]
        text = page.extract_text()
        doc.add_paragraph(text)

    # Ajusta tamaño de fuente (opcional)
    if tamaño_fuente:
        for parrafo in doc.paragraphs:
            parrafo.style.font.size = Pt(tamaño_fuente)

    doc.save(f"{nombre_salida}.docx")

def convertir_word_a_pdf(archivo_word, nombre_salida):
    """Convierte un archivo Word a PDF."""
    doc = Document(archivo_word)
    output_pdf = BytesIO()
    doc.save(output_pdf)
    output_pdf.seek(0)
    writer = PdfWriter()
    writer.add_page(PdfReader(output_pdf).pages[0])
    with open(f"{nombre_salida}.pdf", "wb") as f:
        writer.write(f)

# -------------------------------------------------------------------------------------------------------------------- #

# -------------------------------------------------------------------------------------------------------------------- #
# Funciones para la Interacción con Telegram
# -------------------------------------------------------------------------------------------------------------------- #
@app.on_message(filters.command("start"))
async def start(client, message):
    await message.reply_text(
        f"¡Hola! 👋 Soy un bot que puede combinar archivos PDF, Word y de texto.\n\n"
        f"¡Usa los siguientes comandos para interactuar conmigo!\n"
        f"• `/start`: Iniciar el bot.\n"
        f"• `/help`: Mostrar la ayuda.\n"
        f"• `/enviar`: Enviar archivos para combinar.",
        reply_markup=InlineKeyboardMarkup(
            [
                [InlineKeyboardButton("Enviar Archivo 📁", callback_data="enviar_archivo")],
                [InlineKeyboardButton("Personalizar ⚙️", callback_data="personalizar")]
            ]
        )
    )

@app.on_message(filters.command("help"))
async def help(client, message):
    await message.reply_text(
        f"Comandos disponibles:\n"
        f"• `/start`: Iniciar el bot.\n"
        f"• `/help`: Mostrar la ayuda.\n"
        f"• `/enviar`: Enviar archivos para combinar.",
        reply_markup=InlineKeyboardMarkup(
            [
                [InlineKeyboardButton("Enviar Archivo 📁", callback_data="enviar_archivo")],
                [InlineKeyboardButton("Personalizar ⚙️", callback_data="personalizar")]
            ]
        )
    )

@app.on_message(filters.command("enviar"))
async def enviar_archivo(client, message):
    global documents, current_file_type
    await message.reply_text("¡Envíame un documento o texto! 📝")

@app.on_message(filters.document | filters.text, func=lambda m: m.chat.id == message.chat.id)
async def process_document(client, message):
    global documents, current_file_type
    try:
        if message.document:
            # Obtener el archivo
            file_info = await message.download(file_name=f"{message.document.file_name}")

            # Determinar el tipo de archivo
            if file_info.lower().endswith(('.pdf', '.docx', '.doc', '.txt')):
                current_file_type = file_info.lower().split('.')[-1]
                documents.append(file_info)
                await message.reply_text(f"¡Archivo {file_info} recibido! 🎉 ¿Deseas agregar más archivos del mismo tipo?", reply_markup=InlineKeyboardMarkup(
                    [
                        [InlineKeyboardButton("Sí 👍", callback_data="agregar_mas")],
                        [InlineKeyboardButton("No 👎", callback_data="generar_archivo")]
                    ]
                ))
            else:
                await message.reply_text("¡Solo se aceptan archivos PDF, Word, TXT o texto! ⛔")
                await message.reply_text("¿Deseas enviar otro archivo?", reply_markup=InlineKeyboardMarkup(
                    [
                        [InlineKeyboardButton("Sí 👍", callback_data="enviar_archivo")],
                        [InlineKeyboardButton("No 👎", callback_data="cancelar")]
                    ]
                ))

        elif message.text:
            documents.append(message.text)
            current_file_type = 'text'
            await message.reply_text("¿Deseas agregar más archivos o textos?", reply_markup=InlineKeyboardMarkup(
                [
                    [InlineKeyboardButton("Sí 👍", callback_data="agregar_mas")],
                    [InlineKeyboardButton("No 👎", callback_data="generar_archivo")]
                ]
            ))
    except Exception as e:
        print(f"Error al procesar el archivo: {e}")
        await message.reply_text(f"¡Ups! 😨 Hubo un error al procesar el archivo. Intenta de nuevo.")

@app.on_callback_query(filters.regex(r"^(agregar_mas|generar_archivo|cambiar_formato|pdf|word|personalizar|nombre_archivo|tamaño_fuente|regresar|cancelar)$"))
async def handle_callback_query(client, callback_query):
    global documents, current_file_type, customization_options

    if callback_query.data == "agregar_mas":
        if current_file_type:
            await callback_query.message.reply_text(f"¡Envíame otro archivo de tipo {current_file_type}! 📂")
        else:
            await callback_query.message.reply_text("¡Primero debes enviar un archivo o texto! 📝")
        await callback_query.answer()

    elif callback_query.data == "generar_archivo":
        if current_file_type:
            await callback_query.message.reply_text("¡Elige el formato para el archivo combinado! 🎨", reply_markup=InlineKeyboardMarkup(
                [
                    [InlineKeyboardButton("PDF 📄", callback_data="pdf")],
                    [InlineKeyboardButton("Word 📑", callback_data="word")],
                    [InlineKeyboardButton("Cambiar Formato 🔁", callback_data="cambiar_formato")]
                ]
            ))
        else:
            await callback_query.message.reply_text("¡Primero debes enviar un archivo o texto! 📝")
        await callback_query.answer()

    elif callback_query.data == "pdf":
        await generar_pdf(client, callback_query.message)
        await callback_query.answer()

    elif callback_query.data == "word":
        await generar_word(client, callback_query.message)
        await callback_query.answer()

    elif callback_query.data == "cambiar_formato":
        if current_file_type == 'pdf':
            customization_options['formato'] = 'word'
            await callback_query.message.reply_text("¡El formato se cambiará a Word! 📑")
            await generar_archivo(client, callback_query.message)
        elif current_file_type == 'docx' or current_file_type == 'doc':
            customization_options['formato'] = 'pdf'
            await callback_query.message.reply_text("¡El formato se cambiará a PDF! 📄")
            await generar_archivo(client, callback_query.message)
        else:
            await callback_query.message.reply_text("¡No se ha seleccionado ningún archivo para cambiar el formato! 😕")
        await callback_query.answer()

    elif callback_query.data == "personalizar":
        await callback_query.message.reply_text("¡Elige una opción para personalizar tu archivo combinado! 🎨", reply_markup=InlineKeyboardMarkup(
            [
                [InlineKeyboardButton("Nombre del Archivo 📝", callback_data="nombre_archivo")],
                [InlineKeyboardButton("Tamaño de Fuente 📏", callback_data="tamaño_fuente")],
                [InlineKeyboardButton("Regresar 🔙", callback_data="regresar")]
            ]
        ))
        await callback_query.answer()

    elif callback_query.data == "nombre_archivo":
        await callback_query.message.reply_text("¡Ingresa el nombre que quieres para tu archivo! ✍️")
        @app.on_message(filters.text, func=lambda m: m.chat.id == callback_query.message.chat.id)
        async def set_file_name_input(client, message):
            global customization_options
            customization_options['nombre_archivo'] = message.text
            await message.reply_text(f"¡Nombre del archivo actualizado a: {customization_options['nombre_archivo']}! ✅")
            await handle_personalization(client, message)
        await callback_query.answer()

    elif callback_query.data == "tamaño_fuente":
        await callback_query.message.reply_text("¡Ingresa el tamaño de fuente que quieres (número entero)! 🔢")
        @app.on_message(filters.text, func=lambda m: m.chat.id == callback_query.message.chat.id)
        async def set_font_size_input(client, message):
            global customization_options
            try:
                font_size = int(message.text)
                customization_options['tamaño_fuente'] = font_size
                await message.reply_text(f"¡Tamaño de fuente actualizado a: {customization_options['tamaño_fuente']}! ✅")
                await handle_personalization(client, message)
            except ValueError:
                await message.reply_text("¡Por favor, ingresa un número entero válido! ⛔")
        await callback_query.answer()

    elif callback_query.data == "regresar":
        await callback_query.message.reply_text("¡Regresando al menú principal! 🔙")
        await start(client, callback_query.message)
        await callback_query.answer()

    elif callback_query.data == "cancelar":
        global documents, current_file_type
        documents.clear()
        current_file_type = None
        await callback_query.message.reply_text("¡Proceso cancelado! 🚫")
        await start(client, callback_query.message)
        await callback_query.answer()

async def generar_pdf(client, message):
    global documents, customization_options
    await message.reply_text("¡Procesando tu archivo! ⏳")
    if documents:
        try:
            # Combinar archivos PDF
            if current_file_type == 'pdf':
                output_pdf = BytesIO()
                merger = PdfMerger()
                for archivo in documents:
                    merger.append(archivo)
                merger.write(output_pdf)
                merger.close()

                # Enviar el archivo PDF
                await client.send_document(message.chat.id, output_pdf, filename=f"{customization_options['nombre_archivo']}.pdf", caption="¡Tu archivo combinado está listo! 🎉")

            # Convertir archivos Word a PDF y combinar
            elif current_file_type == 'docx' or current_file_type == 'doc':
                output_pdf = BytesIO()
                merger = PdfMerger()
                for archivo in documents:
                    convertir_word_a_pdf(archivo, customization_options['nombre_archivo'])
                    merger.append(f"{customization_options['nombre_archivo']}.pdf")
                merger.write(output_pdf)
                merger.close()
                # Borrar archivo temporal
                os.remove(f"{customization_options['nombre_archivo']}.pdf")

                # Enviar el archivo PDF
                await client.send_document(message.chat.id, output_pdf, filename=f"{customization_options['nombre_archivo']}.pdf", caption="¡Tu archivo combinado está listo! 🎉")

            # Combinar archivos de texto
            elif current_file_type == 'text':
                output_pdf = BytesIO()
                merger = PdfMerger()
                for archivo in documents:
                    merger.append(archivo)
                merger.write(output_pdf)
                merger.close()

                # Enviar el archivo PDF
                await client.send_document(message.chat.id, output_pdf, filename=f"{customization_options['nombre_archivo']}.pdf", caption="¡Tu archivo combinado está listo! 🎉")

            # Limpiar la lista de documentos
            documents.clear()
            current_file_type = None
        except Exception as e:
            print(f"Error al generar el archivo: {e}")
            await message.reply_text(f"¡Ups! 😨 Hubo un error al generar el archivo. Intenta de nuevo.")
    else:
        await message.reply_text("¡No se han enviado archivos aún! 😕")

async def generar_word(client, message):
    global documents, customization_options
    await message.reply_text("¡Procesando tu archivo! ⏳")
    if documents:
        try:
            # Combinar archivos Word
            if current_file_type == 'docx' or current_file_type == 'doc':
                doc = Document()
                for archivo in documents:
                    doc_temp = Document(archivo)
                    for parrafo in doc_temp.paragraphs:
                        doc.add_paragraph(parrafo.text)
                    # Ajusta tamaño de fuente (opcional)
                    if 'tamaño_fuente' in customization_options:
                        for parrafo in doc.paragraphs:
                            parrafo.style.font.size = Pt(customization_options['tamaño_fuente'])
                output_word = BytesIO()
                doc.save(output_word)
                await client.send_document(message.chat.id, output_word, filename=f"{customization_options['nombre_archivo']}.docx", caption="¡Tu archivo combinado está listo! 🎉")

            # Convertir archivos PDF a Word y combinar
            elif current_file_type == 'pdf':
                doc = Document()
                for archivo in documents:
                    convertir_pdf_a_word(archivo, customization_options['nombre_archivo'], customization_options['tamaño_fuente'])
                    doc_temp = Document(f"{customization_options['nombre_archivo']}.docx")
                    for parrafo in doc_temp.paragraphs:
                        doc.add_paragraph(parrafo.text)
                    # Borrar archivo temporal
                    os.remove(f"{customization_options['nombre_archivo']}.docx")

                output_word = BytesIO()
                doc.save(output_word)
                await client.send_document(message.chat.id, output_word, filename=f"{customization_options['nombre_archivo']}.docx", caption="¡Tu archivo combinado está listo! 🎉")

            # Combinar archivos de texto
            elif current_file_type == 'text':
                doc = Document()
                for archivo in documents:
                    doc.add_paragraph(archivo)
                    # Ajusta tamaño de fuente (opcional)
                    if 'tamaño_fuente' in customization_options:
                        for parrafo in doc.paragraphs:
                            parrafo.style.font.size = Pt(customization_options['tamaño_fuente'])
                output_word = BytesIO()
                doc.save(output_word)
                await client.send_document(message.chat.id, output_word, filename=f"{customization_options['nombre_archivo']}.docx", caption="¡Tu archivo combinado está listo! 🎉")

            # Limpiar la lista de documentos
            documents.clear()
            current_file_type = None
        except Exception as e:
            print(f"Error al generar el archivo: {e}")
            await message.reply_text(f"¡Ups! 😨 Hubo un error al generar el archivo. Intenta de nuevo.")
    else:
        await message.reply_text("¡No se han enviado archivos aún! 😕")


print("Bot iniciado...")
app.run()
